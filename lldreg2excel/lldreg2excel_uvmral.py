# Description: This is a table in word2excel_UVMRAL Python script.
# Author:    huangkk
# Date:       26/01/2022
# Version：   v1.0
# -*- coding : utf-8-*-
# coding:unicode_escape

from __future__ import print_function
import docx
import xlwings as xw
import numpy as np
import re
import pandas as pd
import csv
import os
import sys

# ********************************** 0. 将输出的信息导出到log中 ******************************** #
class Logger:
    def __init__(self, filename="log.txt"):
        self.terminal = sys.stdout
        self.log = open(filename, "w", encoding='utf-8')

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)
        self.log.flush()  # 缓冲区的内容及时更新到log文件中

    def flush(self):
        pass


path = os.path.abspath(os.path.dirname(__file__))
type = sys.getfilesystemencoding()

# ************************************************************************************************************** #
# 函数: word2reg
# 功能：实现将IP UM文档（Secret_IP_IPname_Vxx_LLD.docx）中“4-1节的寄存器说明”生成UVM regmodel所需的excel表格
# ************************************************************************************************************** #
def word2reg():
    # ********************************** 1. 提取word UV文档寄存器说明内容到excel中 ******************************** #
    # Use a breakpoint in the code line below to debug your script.
    # 读取word文件
    # IP_path = 'Secret_IP_IOM_V1.0_LLD.docx'
    IP_path = sys.argv[1]
    doc = docx.Document(IP_path)
    doc_name = re.findall(r"Secret_IP_\S+\.(?:docx|doc)", IP_path)
    IPname = re.findall(r"Secret_IP_(.+?)_\S+", IP_path)    # \S+ 匹配至少任意一个非空字符

    ip_name = str(IPname)[2:-2]
    log_name = 'regmodel_' + ip_name + '.log'
    sys.stdout = Logger(log_name)

    print('Starting extract \"4 寄存器\" within docx of \"{0}\" to excel...'.format(IP_path))
    print('ip_name: ', ip_name)


    # doc = docx.Document(r"C:\Users\huangkk\PycharmProjects\pythonProject\test1.docx")
    # 读取word中的表格，定位到“表4-1 寄存器列表”表格
    table = doc.tables
    reglst_tab_cnt = 0  # 保存寄存器列表的位置数
    for table_cnt in range(len(table)):  # 遍历word中所有表格(table)
        reglst_tab_cnt = table_cnt
        if (table[table_cnt].rows[0].cells[0].text + table[table_cnt].rows[0].cells[1].text) == '寄存器名称偏移地址':  # 当遍历到“表4-1 寄存器列表”下方表格时，退出所有循环
            break
    # print(len(table[reglst_tab_cnt].rows))
    # 将word中读取的数据写入到excel表中
    # 1. 创建一个新的excel表格
    app = xw.App(visible=False, add_book=False)
    app.display_alerts = False  # 关闭一些提示信息，加快运行速度。 默认为True
    app.screen_updating = False  # 关闭更新显示工作表的内容，加快运行速度，默认为True
    # 工作簿
    wb = app.books.add()
    # 工作表
    sht = wb.sheets["sheet1"]
    sht.name = IPname[0] + '_ip'
    # 变量初始化
    row = []
    row_tmp = []
    row0_array = []
    row_tmp_array = []
    row_rst_value = []
    row_cnt_array = []
    row_bit_field_array = []
    rst_val_bin2dci_list_all = []
    reg_cnt = 1
    row0_ary_cnt = 0  # word中每个table行数遍历
    sht_row_cnt = 'a' + str(9)  # excel表格初始a列行坐标
    sht_row_sum = 9  # excel表格a列行数累加，初始值为2
    sht_col_cnt = 'e' + str(9)  # excel表格初始e列行坐标
    sht_col_sum = 9  # excel表格d列行数累加，初始值为2
    sht_col_d = 'd' + str(9)
    sht_row_d_sum = 9

    # 提出简单描述的寄存器表格本文到row0中
    row0 = []
    head_cnt1 = 0   # 计数寄存器列表的表头数量
    for row0_cnt in range(len(table[reglst_tab_cnt].rows)):
        if (table[reglst_tab_cnt].cell(row0_cnt, 0).text + table[reglst_tab_cnt].cell(row0_cnt, 1).text) == '寄存器名称偏移地址':
            head_cnt1 += 1
            print('row0_cnt1: ', row0_cnt)
            continue
        print('row0_cnt: ', row0_cnt)
        
        for i in range(0, 2, 1):  # 提取每一行的name和offset addr
            cell0 = table[reglst_tab_cnt].cell(row0_cnt, i)
            str_arr = cell0.text.split('+')     # 不识别BASE_ADDR时，注释掉78~81行即可
            if i == 1:
                cell0.text = cell0.text.replace('BASE_ADDR+', '')
                print('offset_addr_t: ', cell0.text)
                row0.append(cell0.text)
            elif i == 0:
                row0.append(cell0.text)
        cell0 = table[reglst_tab_cnt].cell(row0_cnt, 5)  # 提取第5行的reset value
        row0.append(cell0.text)
        cell0 = table[reglst_tab_cnt].cell(row0_cnt, 4)  # 提取第4行的访问权限
        row0.append(cell0.text)

    print('row0: ', row0)
    print(len(table[reglst_tab_cnt].rows))
    print(head_cnt1)
    reg_num = len(table[reglst_tab_cnt].rows) - head_cnt1
    print('reg_number: ', reg_num)
    row0_array = np.array(row0).reshape(len(table[reglst_tab_cnt].rows) - head_cnt1, 4)  # 将row0转化成10行3列的数组
    row0_array[:, [0, 1]] = row0_array[:, [1, 0]]  # 交换offset列和name列
    row0_array = np.insert(row0_array, 2, 32, axis=1)  # 在a3列添加'width 32'
    print('寄存器列表：')
    print(row0_array)
    print('寄存器总数：', len(row0_array))
    print()

    # 提出寄存器详细描述表格中文本到row中，并且将“寄存器列表表”表格和“寄存器详细描述表格”按格式添加到excel中
    print('寄存器说明：')
    for table_cnt in range(reglst_tab_cnt + 1, reglst_tab_cnt + len(table[reglst_tab_cnt].rows) - head_cnt1 + 1):  # 遍历word中所有table(表格)
        head_cnt = 0
        for row_cnt in range(0, len(table[table_cnt].rows), 1):  # 遍历table中的所有row(行)
            if (table[table_cnt].cell(row_cnt, 0).text + table[table_cnt].cell(row_cnt, 1).text + table[table_cnt].cell(row_cnt, 2).text) == '位域名称属性':     # 跳过表头
                print(table[table_cnt].cell(row_cnt, 0).text)
                head_cnt += 1
                continue
            for i in range(0, 4, 1):
                cell = table[table_cnt].cell(row_cnt, i)  # 遍历row每一个cell(单元格)
                row.append(cell.text)  # 将cell中的文本追加到row数组中
        row_array = np.array(row).reshape(len(table[table_cnt].rows) - head_cnt, 4)  # 将row数组转化成10行4列的10维数组
        # row_array中除去一个数组(表头)的其他所有数组添加到row_tmp中，并且将每个数组中的bits变量字符串左右两边拼接'['和']'
        
        i = -1
        reserved_cnt = 0
        reserved_arr = [
            'Reserved0', 'Reserved1', 'Reserved2', 'Reserved3', 'Reserved4', 'Reserved5', 'Reserved6', 'Reserved7',
            'Reserved8', 'Reserved9', 'Reserved10', 'Reserved11', 'Reserved12', 'Reserved13', 'Reserved14', 'Reserved15',
            'Reserved16', 'Reserved17', 'Reserved18', 'Reserved19', 'Reserved20', 'Reserved21', 'Reserved22', 'Reserved23',
            'Reserved24', 'Reserved25', 'Reserved26', 'Reserved27', 'Reserved28', 'Reserved29', 'Reserved30', 'Reserved31',
        ]
        
        for row_cnt in range(len(row_array)):
            row_tmp.append('[' + row_array[row_cnt][0] + ']')
            for ele in row_array[row_cnt][1:]:
                row_tmp.append(ele)

        for row_reg in reversed(row_tmp):   # 将位域中名称为‘-’替换为‘Reservedx’（x为自然数）
            if row_reg == '-':
                row_reg = reserved_arr[reserved_cnt]
                row_tmp[i] = row_reg
                reserved_cnt += 1
                print('row_reg: ', row_reg)
            i -= 1
        print('row_tmp: ', row_tmp)

        row_tmp_array = np.array(row_tmp).reshape(len(table[table_cnt].rows) - head_cnt, 4)
        row_tmp1_array = row_tmp_array[::-1]    # 反转
        print('num: ', reg_cnt)
        print(row_tmp_array)

        # ******************************************* REGm ******************************************************* #
        # print(row0_array[row0_ary_cnt:row0_ary_cnt + 1])
        # print(row0_array[row0_ary_cnt][0])
        # 定位多个相同功能的寄存器所在位置
        if '+' in row0_array[row0_ary_cnt][0]:      # 例如：0x3C+i*0x4(i=0:15)
            reg_name = row0_array[row0_ary_cnt][1]

            offset_addr_list = row0_array[row0_ary_cnt][0].split('+')

            # 1. 提取第一个寄存器的地址，即 0x3C
            offset_addr_first = offset_addr_list[0]     # first offset addr

            # 2. 提取偏移量offset，即 0x4
            offset_list1 = offset_addr_list[1].split('(')
            offset_tmp1 = offset_list1[0]
            offset_tmp2 = offset_tmp1.split('*')
            offset = offset_tmp2[1]     # offset = 0x4

            # 提取相同功能寄存器的总数，即 15 + 1 = 16
            offset_addr_list_str = list(row0_array[row0_ary_cnt][0])
            print(offset_addr_list_str)
            for chr in range(len(offset_addr_list_str)):
                if offset_addr_list_str[chr] == '=':
                    offset_addr_num = int(''.join(offset_addr_list_str[chr+1:-1]))   # reg number
                    print(''.join(offset_addr_list_str[chr+1:-1]))
                    break
            print(offset_addr_num)
            for i in range(offset_addr_num):
                # print(i)
                offset_addr_tmp = hex(eval(offset_addr_first) + i * eval(offset))
                offset_addr = '0x' + offset_addr_tmp[2:].upper()
                row0_array[row0_ary_cnt][0] = offset_addr
                row0_array[row0_ary_cnt][1] = row0_array[row0_ary_cnt][1] + str(i)
                sht.range(sht_row_cnt).value = row0_array[row0_ary_cnt:row0_ary_cnt + 1]  # 将row0_array数组中从index从0开始的每一行添加到sht_row_cnt行中
                sht.range(sht_col_cnt).value = row_tmp1_array  # 将row_tmp_array数组每一行添加到sht_col_cnt行中
                sht_row_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_row_sum累加每个表格的行数
                sht_col_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_col_sum累加每个表格的行数
                sht_row_cnt = 'a' + str(sht_row_sum)  # 'a' + str(sht_row_sum - 1)转换成a列单元格坐标
                sht_col_cnt = 'e' + str(sht_col_sum)  # 'e' + str(sht_col_sum1 - 1)转换成e列单元格坐标
                row0_array[row0_ary_cnt][1] = reg_name

                # ************************************** 获取所需变量值 ********************************************* #
                # 存储每个寄存器的行数
                row_cnt_array.append(len(table[table_cnt].rows) - 1)

                # 将每个寄存器中的Bit组成一个列表，所有寄存器的每个Bit列表组成一个多维列表，函数rst_val_bin2dci_list_all需要
                row_tmp_array1 = []
                for row_cnt in range(len(row_tmp_array)):
                    row_tmp_array1.append(row_tmp_array[row_cnt][0])
                row_bit_field_array.append(row_tmp_array1)

                # 存储每个寄存器的复位值
                row_rst_value.append(row0_array[row0_ary_cnt][3])

                # print(offset_addr)
                # print(hex(eval(offset_addr_list[0]) + i * 0x4))
            row = []  # row数组初始化，继续添加添加下一个表格文本
            row_tmp = []  # row_tmp数组初始化，继续添加下一个row_array数组
            row0_ary_cnt += 1  # row0_array数组行数加1
            reg_cnt += 1  # 累加寄存器数量
        # ******************************************** REGm done ****************************************************** #

        else:
            sht.range(sht_row_cnt).value = row0_array[row0_ary_cnt:row0_ary_cnt + 1]    # 将row0_array数组中从index从0开始的每一行添加到sht_row_cnt行中
            sht.range(sht_col_cnt).value = row_tmp1_array  # 将row_tmp_array数组每一行添加到sht_col_cnt行中

            sht_row_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_row_sum累加每个表格的行数
            sht_col_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_col_sum累加每个表格的行数
            sht_row_cnt = 'a' + str(sht_row_sum)  # 'a' + str(sht_row_sum - 1)转换成a列单元格坐标
            sht_col_cnt = 'e' + str(sht_col_sum)  # 'e' + str(sht_col_sum1 - 1)转换成e列单元格坐标

            # ************************************** 获取所需变量值 ********************************************* #
            # 获取每一个详细列表除去表头的行数，函数cell_merge_col_ABC需要
            row_cnt_array.append(len(table[table_cnt].rows) - head_cnt)

            # 将每个寄存器中的Bit组成一个列表，所有寄存器的每个Bit列表组成一个多维列表，函数rst_val_bin2dci_list_all需要
            row_tmp_array1 = []
            for row_cnt in range(len(row_tmp_array)):
                row_tmp_array1.append(row_tmp_array[row_cnt][0])
            row_bit_field_array.append(row_tmp_array1)

            # 存储每个寄存器的复位值
            row_rst_value.append(row0_array[row0_ary_cnt][3])

            row = []  # row数组初始化，继续添加添加下一个表格文本
            row_tmp = []  # row_tmp数组初始化，继续添加下一个row_array数组
            row0_ary_cnt += 1  # row0_array数组行数加1
            reg_cnt += 1    # 累加寄存器数量

        # *********************************************** 1. done ********************************************* #
        print()
    print(row_rst_value)

    # ****************************************** 2. done ***************************************************** #

    # ********************************** 3.  调用函数：cell_merge_col_ABC ************************************* #
    cell_merge_col_ABC(sht, row_cnt_array)

    # ********************************** 4. 调用函数：rst_val2dci_val  **************************************** #
    rst_val2dci_val(sht, row_rst_value, row_bit_field_array, row_cnt_array, sht_col_d, sht_row_d_sum)

    # ********************************** 5. 调用函数： modify_cell_format ************************************* #
    modify_cell_format(sht)

    # ********************************** 6. 调用函数： exchange_col ******************************************* #
    exchange_col(sht)

    # ********************************** 7. 保存excel表格并关闭xwlings程序 ************************************* #
    # 保存excel
    reg_filename_xlsx = 'regmodel_' + IPname[0] + '.xlsx'
    wb.save(reg_filename_xlsx)
    # 关闭xwlings程序
    wb.close()
    app.quit()


# ******************************************************************** #
# 函数：modify_cell_format
# 功能：更改1~7行单元格颜色，合并单元格以及为单元格边框增加边框线
# ******************************************************************** #
def modify_cell_format(sht):
    sht.range('a1').value = ['ProjectName']
    sht.range('a1:h1').color = (124, 252, 0)
    sht.range('b1:h1').merge()
    add_border(sht, 'a1')
    add_border(sht, 'b1:h1')

    sht.range('a2').value = ['Module']
    sht.range('a2:h2').color = (124, 252, 0)
    sht.range('b2:h2').merge()
    add_border(sht, 'a2')
    add_border(sht, 'b2:h2')

    sht.range('a3').value = ['Protocal']
    sht.range('a3:h3').color = (124, 252, 0)
    sht.range('b3:h3').merge()
    add_border(sht, 'a3')
    add_border(sht, 'b3:h3')

    sht.range('a4').value = ['BaseADDR']
    sht.range('a4:h4').color = (124, 252, 0)
    sht.range('b4:h4').merge()
    add_border(sht, 'a4')
    add_border(sht, 'b4:h4')

    sht.range('a5').value = ['AddrWidth']
    sht.range('a5:h5').color = (124, 252, 0)
    sht.range('b5:h5').merge()
    add_border(sht, 'a5')
    add_border(sht, 'b4:h5')

    sht.range('a6').value = ['Version']
    sht.range('a6:h6').color = (124, 252, 0)
    sht.range('b6:h6').merge()
    add_border(sht, 'a6')
    add_border(sht, 'b6:h6')

    sht.range('a7').value = ['#REGISTER_DEFINE#']
    sht.range('a7:h7').color = (192, 192, 192)

    sht.range('a8').value = ['Offset', 'RegName', 'Width', 'Reset', 'Bit', 'Field', 'Access',
                             'Filed Description']  # 将表头添加到a1单元格所在行


# ******************************************************************** #
# 函数：exchange_col
# 功能：将第5，6，7列整体左移一列，将第4列复制到第9列，然后复制到第7列
# ******************************************************************** #
def exchange_col(sht):
    sht.api.Columns(4).Copy(sht.api.Columns(9))  # 将第4列复制到第9列
    sht.api.Columns(5).Copy(sht.api.Columns(4))  # 第5列复制到第4列
    sht.api.Columns(6).Copy(sht.api.Columns(5))  # ...
    sht.api.Columns(7).Copy(sht.api.Columns(6))  # ...
    sht.api.Columns(9).Copy(sht.api.Columns(7))  # ...
    sht.range('i7').api.EntireColumn.Delete()  # 将i7列删除
    sht.autofit()  # 按文本内容自动扩充单元格


# ******************************************************************** #
# 函数：add_border
# 功能：单元格添加边框
# ******************************************************************** #
def add_border(sht, cell):
    for border_id in range(7, 11):
        sht.range(cell).api.Borders(border_id).LineStyle = 1


# ******************************************************************** #
# 函数：cell_merge_col_ABC
# 功能：合并A, B, C列单元格
# ******************************************************************** #
def cell_merge_col_ABC(sht, row_cnt_array):
    sht_row_a = 'a'
    sht_row_b = 'b'
    sht_row_c = 'c'
    sht_row_sum = 9
    for i in range(len(row_cnt_array)):
        merge_range_a = sht_row_a + str(sht_row_sum) + ':' + sht_row_a + str(sht_row_sum + row_cnt_array[i] - 1)
        merge_range_b = sht_row_b + str(sht_row_sum) + ':' + sht_row_b + str(sht_row_sum + row_cnt_array[i] - 1)
        merge_range_c = sht_row_c + str(sht_row_sum) + ':' + sht_row_c + str(sht_row_sum + row_cnt_array[i] - 1)
        sht.range(merge_range_a).merge()
        sht.range(merge_range_b).merge()
        sht.range(merge_range_c).merge()
        sht_row_sum += row_cnt_array[i]
        # print(merge_range_a)
        # print(merge_range_b)
        # print(merge_range_c)


# ******************************************************************** #
# 函数：rst_val2dci_val
# 功能：将复位值的每个bit_field拆分出来并转换成10进制数(核心函数)
# ******************************************************************** #
def rst_val2dci_val(sht, row_rst_value, row_bit_field_array, row_cnt_array, sht_col_d, sht_row_d_sum):
    rst_val_bin2dci_list_all_tmp = []
    print('hex to bin: ')
    for i in range(len(row_rst_value)):
        rst_val_bin2dci_list = []
        rst_val_hex = row_rst_value[i]  # 给定一个32位的16进制数
        rst_val_bin = bin(int(rst_val_hex, 16))  # 将该16进制转化成2进制数，会自动丢弃从从低到高位第一个为0及其之后全为0的比特位
        rst_val_bin_list_tmp = list(rst_val_bin[2:])  # 切除该2进制数左边的2进制数标记‘0b’
        rst_val_bin_list_str = (''.join(rst_val_bin_list_tmp)).zfill(32)  # 将该2进制数填充自动丢弃的0为32位
        rst_val_bin_list = list(rst_val_bin_list_str)
        print('num: ', i + 1)
        print(rst_val_bin_list)

        bit_field_list = []
        bit_field = row_bit_field_array[i]
        dci_field = []
        dci_arr1 = []
        dci_arr2 = []
        bit_field_sum = 0

        # 计算每个寄存每个bit field的比特总数
        for i in range(len(bit_field)):
            bit_field_list = list(bit_field[i].strip('[').strip(']'))
            if ':' in bit_field_list:
                for j in range(len(bit_field_list)):
                    if bit_field_list[j] == ':':
                        dci_arr1 = ''.join(bit_field_list[:j])
                        dci_arr2 = ''.join(bit_field_list[j + 1:])
                        bit_field_sum = int(dci_arr1) - int(dci_arr2) + 1
                        # print(dci_arr1)
                        # print(dci_arr2)
                        # print()
                    else:
                        continue
            else:
                bit_field_sum = 1
            dci_field.append(bit_field_sum)

        # print(dci_field)

        j = 0
        for i in range(len(dci_field)):
            # print(dci_field[i])
            if j == 0:
                rst_val_bin2str_tmp = rst_val_bin_list[0:dci_field[i]]
            else:
                rst_val_bin2str_tmp = rst_val_bin_list[j:j + dci_field[i]]
            rst_val_bin2str = ''.join(rst_val_bin2str_tmp)
            rst_val_bin2dci_list.append(int(rst_val_bin2str, 2))
            # print(rst_val_bin2str)
            # print(rst_val_bin2dci_list)
            # rst_val_bin2dci_list.append(int(rst_val_bin2str, 2))
            j = j + dci_field[i]
        rst_val_bin2dci_list_all_tmp.append(rst_val_bin2dci_list)
    print()

    for row_cnt in range(0, len(rst_val_bin2dci_list_all_tmp)):
        rst_val_bin2dci_list_all_tmp[row_cnt].reverse()
        sht.range(sht_col_d).options(transpose=True).value = rst_val_bin2dci_list_all_tmp[row_cnt]
        sht_row_d_sum += row_cnt_array[row_cnt]
        sht_col_d = 'd' + str(sht_row_d_sum)

    print('bit field to dcimal: ')
    for i in range(len(rst_val_bin2dci_list_all_tmp)):
        print('num: ', i + 1)
        print(rst_val_bin2dci_list_all_tmp[i])
    print()

    print('Extracted Completed!!!')


# ****************************************** main 函数 **************************************** #
if __name__ == '__main__':
    word2reg()
    input('Press <Enter> to exit ......')

