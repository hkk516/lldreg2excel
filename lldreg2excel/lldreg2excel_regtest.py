# This is a table in word2excel Python script.
# Author:    huangkk
# Date:       12/02/2021
# Version：   v1.0
# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
# -*- coding : utf-8-*-
# coding:unicode_escape
from __future__ import print_function

import sys

import docx
import xlwings as xw
import numpy as np
import re
import pandas as pd
import csv
import os
import copy

# ********************************** 0. 将输出的信息导出到log中 ******************************** #
class Logger:
    def __init__(self, filename="log.txt"):
        self.terminal = sys.stdout
        self.log = open(filename, "w", encoding='utf-8')

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)
        self.log.flush()  # 缓冲区的内容及时更新到log文件中

    def flush(self):
        pass


path = os.path.abspath(os.path.dirname(__file__))
type = sys.getfilesystemencoding()


def word2reg():
    # Use a breakpoint in the code line below to debug your script.
    # 读取word文件
    # IP_path = 'Secret_IP_IOM_V1.0_LLD.docx'    # python无法读取.doc后缀的word，需要将*.doc另存为*.docx
    IP_path = sys.argv[1]
    doc = docx.Document(IP_path)
    # doc_name = re.findall(r"Secret_IP_\S+\.(?:docx|doc)", IP_path)
    IPname = re.findall(r"Secret_IP_(.+?)_\S+", IP_path)    # \S+ 匹配至少任意一个非空字符

    ip_name = str(IPname)[2:-2]
    log_name = 'regfile_' + ip_name + '.log'
    sys.stdout = Logger(log_name)

    print('Starting extract \"4 寄存器\" within docx of \"{0}\" to excel...'.format(IP_path))
    print('ip_name: ', ip_name)

    # doc = docx.Document(r"C:\Users\huangkk\PycharmProjects\pythonProject\test1.docx")
    # 读取word中的表格，定位到“表4-1 寄存器列表”表格
    table = doc.tables
    reglst_tab_cnt = 0  # 保存寄存器列表的位置数
    for table_cnt in range(len(table)):  # 遍历word中所有表格(table)
        reglst_tab_cnt = table_cnt
        if (table[table_cnt].rows[0].cells[0].text + table[table_cnt].rows[0].cells[1].text) == '寄存器名称偏移地址':  # 当遍历到“表4-1 寄存器列表”下方表格时，退出所有循环
            break

    # 将word中读取的数据写入到excel表中
    # 1. 创建一个新的excel表格
    try:
        app = xw.App(visible=False, add_book=False)
        app.display_alerts = False  # 关闭一些提示信息，加快运行速度。 默认为True
        app.screen_updating = False  # 关闭更新显示工作表的内容，加快运行速度，默认为True
        # 工作簿
        wb = app.books.add()
        # 工作表
        sht = wb.sheets["sheet1"]
    except:
        print("Error: can't build a new excel!")
    # 变量初始化
    row = []
    row_tmp = []
    row0_ary_cnt = 0            # word中每个table行数遍历
    sht_row_cnt = 'a' + str(4)  # excel表格初始a列行坐标
    sht_row_sum = 4             # excel表格a列行数累加，初始值为2
    sht_col_cnt = 'e' + str(4)  # excel表格初始e列行坐标
    sht_col_sum = 4             # excel表格d列行数累加，初始值为2
    sht_row_cnt_e_se = 'h' + str(4)     # excel表格初始h列行坐标
    sht_row_cnt_rst_tri = 'i' + str(4)  # excel表格初始化i列坐标

    # 提出简单描述的寄存器表格本文到row0中
    row0 = []
    e_se_protect = []
    rst_tri = []
    head_cnt1 = 0   # 计数寄存器列表的表头数量
    for row0_cnt in range(len(table[reglst_tab_cnt].rows)):
        if (table[reglst_tab_cnt].cell(row0_cnt, 0).text + table[reglst_tab_cnt].cell(row0_cnt, 1).text) == '寄存器名称偏移地址':
            head_cnt1 += 1
            continue
        print('row0_cnt: ', row0_cnt)

        for i in range(7):  # 提取每一行的name和offset addr
            cell0 = table[reglst_tab_cnt].cell(row0_cnt, i)
            str_arr = cell0.text.split('+')     # 不识别BASE_ADDR时，注释掉78~81行即可
            if i == 1:
                cell0.text = cell0.text.replace('BASE_ADDR+', '')
                print('offset_addr_t: ', cell0.text)
                row0.append(cell0.text)
            elif i == 0:
                row0.append(cell0.text)
            elif i == 5:
                row0.append(cell0.text)
            elif i == 4:
                e_se_protect.append(cell0.text)
            elif i == 6:
                rst_tri.append(cell0.text)
        # cell0 = table[reglst_tab_cnt].cell(row0_cnt, 6)  # 提取第5行的reset value
        # row0.append(cell0.text)
        # cell0 = table[reglst_tab_cnt].cell(row0_cnt, 5)  # 提取第4行的访问权限
        # e_se_protect.append(cell0.text)
        #
        # # 在这里改
        # cell0 = table[reglst_tab_cnt].cell(row0_cnt, 7)  # 提取第7行的复位源
        # rst_tri.append(cell0.text)

    cell0 = table[reglst_tab_cnt].cell(2, 5)    # 上面的for循环不知道为什么提取不到第二行的访问权限和复位源，只能单独提取了
    e_se_protect.insert(0, cell0.text)
    cell0 = table[reglst_tab_cnt].cell(2, 7)
    rst_tri.insert(0, cell0.text)

    #e_se_protect = np.delete(e_se_protect, 0, axis=0)
    del e_se_protect[0]
    print('e_se:', e_se_protect)
    #rst_tri = np.delete(rst_tri, 0, axis=0)
    del rst_tri[0]
    print('row0: ', row0)
    print('reset trigger: ', rst_tri)
    print('head_cnt1: ', head_cnt1)
    reg_num = len(table[reglst_tab_cnt].rows) - head_cnt1
    print('reg_number: ', reg_num)

    row0_array = np.array(row0).reshape(len(table[reglst_tab_cnt].rows) - head_cnt1, 3)  # 将row0转化成10行3列的数组
    #row0_array = np.delete(row0_array, 0, axis=0)  # 将row0_array的表头(第1行)删除
    row0_array[:, [0, 1]] = row0_array[:, [1, 0]]  # 交换offset列和name列
    row0_array = np.insert(row0_array, 2, 32, axis=1)   # 在a3列添加'width 32'
    print('row0_array:', row0_array)

    add_ele_num = 0
    tmp_num = 0
    e_se_protect_tmp = copy.deepcopy(e_se_protect)
    rst_tri_tmp = copy.deepcopy(rst_tri)
    for i in range(len(row0_array)):
        if '+' in row0_array[i][0]:
            e_se_local_char = e_se_protect_tmp[i]
            rst_tri_local_char = rst_tri_tmp[i]
            tmp_str_lst = list(row0_array[i][0])
            for ch_num in range(len(tmp_str_lst)):
                if tmp_str_lst[ch_num] == '=':
                    tmp_num = int(''.join(tmp_str_lst[ch_num + 1:-1]))
                    break
            for j in range(1, tmp_num):
                e_se_protect.insert(i + j + add_ele_num, e_se_local_char)
                rst_tri.insert(i + j + add_ele_num, rst_tri_local_char)
            add_ele_num += (tmp_num - 1)

    print('e_se_protect1: ', e_se_protect, len(e_se_protect))
    print('rst_tri1: ', rst_tri, len(rst_tri))

    # 提出寄存器详细描述表格中文本到row中，并且将“寄存器列表表”表格和“寄存器详细描述表格”按格式添加到excel中
    for table_cnt in range(reglst_tab_cnt + 1, reglst_tab_cnt + len(table[reglst_tab_cnt].rows) - head_cnt1 + 1):  # 遍历word中所有table(表格)
        head_cnt = 0
        for row_cnt in range(0, len(table[table_cnt].rows), 1):  # 遍历table中的所有row(行)
            if (table[table_cnt].cell(row_cnt, 0).text + table[table_cnt].cell(row_cnt, 1).text + table[table_cnt].cell(row_cnt, 2).text) == '位域名称属性':     # 跳过表头
                print(table[table_cnt].cell(row_cnt, 0).text)
                head_cnt += 1
                continue
            for i in range(0, 3, 1):
                cell = table[table_cnt].cell(row_cnt, i)  # 遍历row每一个cell(单元格)
                row.append(cell.text)  # 将cell中的文本追加到row数组中
        row_array = np.array(row).reshape(len(table[table_cnt].rows) - head_cnt, 3)  # 将row数组转化成10行3列的10维数组
        # row_array中除去一个数组(表头)的其他所有数组添加到row_tmp中，并且将每个数组中的bits变量字符串左右两边拼接'['和']'

        i = -1
        reserved_cnt = 0
        reserved_arr = [
            'Reserved0', 'Reserved1', 'Reserved2', 'Reserved3', 'Reserved4', 'Reserved5', 'Reserved6', 'Reserved7',
            'Reserved8', 'Reserved9', 'Reserved10', 'Reserved11', 'Reserved12', 'Reserved13', 'Reserved14', 'Reserved15',
            'Reserved16', 'Reserved17', 'Reserved18', 'Reserved19', 'Reserved20', 'Reserved21', 'Reserved22', 'Reserved23',
            'Reserved24', 'Reserved25', 'Reserved26', 'Reserved27', 'Reserved28', 'Reserved29', 'Reserved30', 'Reserved31',
        ]
        for row_cnt in range(0, len(row_array)):
            row_tmp.append('[' + row_array[row_cnt][0] + ']')
            for ele in row_array[row_cnt][1:]:
                row_tmp.append(ele)
        print('row_tmp: ', row_tmp)

        for row_reg in reversed(row_tmp):   # 将位域中名称为‘-’替换为‘Reservedx’（x为自然数）
            # print(row_reg)
            # print(type(row_reg))
            if row_reg == '-':
                row_reg = reserved_arr[reserved_cnt]
                row_tmp[i] = row_reg
                reserved_cnt += 1
                print('row_reg: ', row_reg)
            i -= 1

        row_tmp_array = np.array(row_tmp).reshape(len(table[table_cnt].rows) - head_cnt, 3)
        print('row_tmp_array', row_tmp_array)

        # ******************************************* REGm ******************************************************* #
        # print(row0_array[row0_ary_cnt:row0_ary_cnt + 1])
        # print(row0_array[row0_ary_cnt][0])
        # 定位多个相同功能的寄存器所在位置
        if '+' in row0_array[row0_ary_cnt][0]:      # 例如：0x3C+i*0x4(i=0:15)
            reg_name = row0_array[row0_ary_cnt][1]

            offset_addr_list = row0_array[row0_ary_cnt][0].split('+')
            print('offset_addr_list: ', offset_addr_list)

            # 1. 提取第一个寄存器的地址，即 0x3C
            offset_addr_first = offset_addr_list[0]     # first offset addr
            print('offset_addr_first: ', offset_addr_first)

            # 2. 提取偏移量offset，即 0x4
            offset_list1 = offset_addr_list[1].split('(')
            offset_tmp1 = offset_list1[0]
            offset_tmp2 = offset_tmp1.split('*')
            offset = offset_tmp2[1]     # offset = 0x4

            # 提取相同功能寄存器的总数，即 15 + 1 = 16
            offset_addr_list_str = list(row0_array[row0_ary_cnt][0])
            print(offset_addr_list_str)
            for chr in range(len(offset_addr_list_str)):
                # i=0:7
                # if offset_addr_list_str[chr] == ':':
                #     offset_addr_num = int(''.join(offset_addr_list_str[chr+1:-1])) + 1     # reg number
                #     print(''.join(offset_addr_list_str[chr+1:-1]))
                # i=8
                if offset_addr_list_str[chr] == '=':
                    offset_addr_num = int(''.join(offset_addr_list_str[chr+1:-1]))    # reg number
                    break
            print(offset_addr_num)

            for i in range(offset_addr_num):
                # print(i)
                offset_addr_tmp = hex(eval(offset_addr_first) + i * eval(offset))
                offset_addr = '0x' + offset_addr_tmp[2:].upper()
                row0_array[row0_ary_cnt][0] = offset_addr
                row0_array[row0_ary_cnt][1] = row0_array[row0_ary_cnt][1] + str(i)
                sht.range(sht_row_cnt).value = row0_array[
                                               row0_ary_cnt:row0_ary_cnt + 1]  # 将row0数组中从index从0开始的每一行添加到sht_row_cnt行中
                sht.range(sht_row_cnt_e_se).value = e_se_protect[
                                                    row0_ary_cnt:row0_ary_cnt + 1]
                sht.range(sht_row_cnt_rst_tri).value = rst_tri[
                                                       row0_ary_cnt:row0_ary_cnt + 1]
                sht.range(sht_col_cnt).value = row_tmp_array  # 将row_tmp_array数组每一行添加到sht_col_cnt行中
                sht_row_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_row_sum累加每个表格的行数
                sht_col_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_col_sum1累加每个表格的行数
                sht_row_cnt = 'a' + str(sht_row_sum)  # 'a' + str(sht_row_sum - 1)转换成a列单元格坐标
                sht_col_cnt = 'e' + str(sht_col_sum)  # 'e' + str(sht_col_sum1 - 1)转换成e列单元格坐标
                sht_row_cnt_e_se = 'h' + str(sht_row_sum)
                sht_row_cnt_rst_tri = 'i' + str(sht_row_sum)
                row0_array[row0_ary_cnt][1] = reg_name
            row = []  # row数组初始化，继续添加添加下一个表格文本
            row_tmp = []  # row_tmp数组初始化，继续添加下一个row_array数组
            row0_ary_cnt += 1  # row0数组行数加1
        # ************************ REGm done **************************  #

        else:
            sht.range(sht_row_cnt).value = row0_array[
                                           row0_ary_cnt:row0_ary_cnt + 1]  # 将row0数组中从index从0开始的每一行添加到sht_row_cnt行中
            sht.range(sht_row_cnt_e_se).value = e_se_protect[
                                                row0_ary_cnt:row0_ary_cnt + 1]
            sht.range(sht_row_cnt_rst_tri).value = rst_tri[
                                                   row0_ary_cnt:row0_ary_cnt + 1]
            sht.range(sht_col_cnt).value = row_tmp_array  # 将row_tmp_array数组每一行添加到sht_col_cnt行中
            row = []  # row数组初始化，继续添加添加下一个表格文本
            row_tmp = []  # row_tmp数组初始化，继续添加下一个row_array数组
            row0_ary_cnt += 1  # row0数组行数加1
            sht_row_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_row_sum累加每个表格的行数
            sht_col_sum += (len(table[table_cnt].rows) - head_cnt)  # sht_col_sum1累加每个表格的行数
            sht_row_cnt = 'a' + str(sht_row_sum)  # 'a' + str(sht_row_sum - 1)转换成a列单元格坐标
            sht_col_cnt = 'e' + str(sht_col_sum)  # 'e' + str(sht_col_sum1 - 1)转换成e列单元格坐标
            sht_row_cnt_e_se = 'h' + str(sht_row_sum)
            sht_row_cnt_rst_tri = 'i' + str(sht_row_sum)
    sht.range('a1').value = ['ipnumbers', '']
    sht.range('a2').value = ['baseaddr', '']
    sht.range('a3').value = ['offset', 'name', 'width', 'Reset', 'bit', 'field', 'Access', 'e_se_protect', 'reset_trigger']  # 将表头添加到a1单元格所在行
    # sht.autofit()  # 按文本内容自动扩充单元格
    print('e_se_protect: ', e_se_protect)
    print('rst_tri: ', rst_tri)
    # 保存excel
    reg_filename_xlsx = 'regfile_' + IPname[0] + '.xlsx'
    wb.save(reg_filename_xlsx)
    # 关闭excel程序
    wb.close()
    app.quit()

    """将.xlsx转化为.csv"""
    xlsx_to_csv_pd(reg_filename_xlsx)

    print('start SerchBasAd')
    baseAd=SerchBasAd('regfile.csv')
    regstruct=csv2reg('regfile.csv')
    print('start writefile')
    writefile('regtest.c',baseAd,regstruct)
    baseAd=SerchBasAd('regfile.csv')
    demo_case_transform('regtest.c', row0_array, e_se_protect, rst_tri, head_cnt1)
    # delete_excel_regtest_file(reg_filename_xlsx, 'regtest.c')

    print("Extraction \"{0}\" done!!!".format('regtest.c'))


def xlsx_to_csv_pd(reg_filename_xlsx):
    data_xlsx = pd.read_excel(reg_filename_xlsx, index_col=0)
    data_xlsx.to_csv('regfile.csv', encoding='utf-8')


def demo_case_transform(demo_case_name, row0_array, e_se_protect, rst_tri, head_cnt1):
    # 给IP_Reg结构体数组添加逗号
    num1 = 0    # 定位到'RegTypeDef IP_Reg[]'所在的行
    num2 = 0    # 定位到'RegTypeDef IP_Reg[]'内容的第一行：lines[num2]:  	{ 0x30 -1 0xffffffff },
    num3 = 1
    num4 = 0
    num5 = 0
    num6 = 0
    num7 = 0
    try:
        with open(demo_case_name, "r") as file_obj:
            lines = file_obj.readlines()
    except:
        print("Error: can't open demo case of \"demo_case.c\"")

    try:
        file_obj1 = open('regtest.c', "w")
    except:
        print("Error: can't open file of \"regtest.c\"")

    for num1 in range(len(lines)):
        file_obj1.write(lines[num1])
        if 'RegTypeDef IP_Reg[] =' in lines[num1]:
            num2 = num1 + 3  # 定位到结构体的第一行
            print('lines[num2]: ', lines[num2])
            break
    file_obj1.write(lines[num1 + 1])
    file_obj1.write(lines[num1 + 2])

    # num1 = num1 + 2

    ### 转换的代码
    print('row0_array: ', row0_array)
    for num4 in range(num2, num2 + len(e_se_protect) * 2 - 1):
        str_list = list(lines[num4])  # 字符串转list
        for num5 in range(len(str_list)):           # 添加offset后面的逗号
            if str_list[num5] == ' ':
                if num3 == 2:
                    str_list.insert(num5, ',')
                num3 = num3 + 1

        num3 = 1
        for num5 in range(len(str_list)):           # 添加DefaultValue后面的逗号
            if str_list[num5] == ' ':
                if num3 == 3:
                    str_list.insert(num5, ',')
                num3 = num3 + 1
        print('str_list: ', str_list)

        num3 = 1
        for num5 in range(len(str_list)):           # 在rw后面添加e_se_protect, 在e_se_protect后面添加app_rst_n&sys_rst_n
            if str_list[num5] == ' ':
                if num3 == 4:
                    str_list.insert(num5, ',' + ' ' + '"' + e_se_protect[num6] + '"' + ',' + ' ' + '"' + rst_tri[num6] + '"')
                    num6 = num6 + 1
                num3 = num3 + 1

        str_out = ''.join(str_list)  # 空字符连接
        file_obj1.write(str_out)
        print('str_out', str_out)
        ### 转换完毕

    for str in lines[num4 + 1:]:
        file_obj1.write(str)

# 删除生成的中间文件：excel和regtest.c
def delete_excel_regtest_file(file_excel, file_regtest):
    if os.path.exists(file_excel):
        os.remove(file_excel)
    else:
        print("The file_excel does not exist")

    if os.path.exists(file_regtest):
        os.remove(file_regtest)
    else:
        print("The file_regtest does not exist")

# 从csv中找到baseaddr并将其提取出来, 并将其填入module_baseaddr数组
def SerchBasAd (filename):
    with open (filename,"r", encoding='gbk', errors='ignore') as fileid:
        reader=csv.reader(fileid)
        #for row in reader:
        #    print(row)
        # 将csv中的每一行文本内容写入到flist列表中
        flist=list(reader)
        #print('filelist')
        #print(flist)
        linenum=0
        # 遍历flist列表
        for templist in flist:
            linenum=linenum+1
            print('templist:', templist)
            if "baseaddr" in templist:  # 若baseaddr在该行中，则将其后的baseaddr写入到BaseAd中
                BaseAd=templist[1:]

        # 将baseaddr写入到result列表中
        result=[]
        for RegAd in BaseAd:
            if RegAd!='':
                result.append(RegAd)
                print('result:', result)
            else:
                pass
        return result

# 从csv中找到每个寄存器的offset并将其提取出来, 将其填入结构体数组IP_Reg的offset变量中
def SerchRegAd (filename):
    with open (filename,"r", encoding='gbk', errors='ignore') as fileid:
        reader=csv.reader(fileid)
        #for row in reader:
        #    print(row)
        flist=list(reader)
        #print('filelist')
        #print(flist)
        linenum=0
        RegAd=[]
        for templist in flist:
            linenum=linenum+1
            print(templist)
            if "offset" in templist:
                RegAd.append(linenum)
                print('offset')
                print(RegAd)
        return RegAd


def writefile (codename,baseaddr,regstruct):
    wfid=open(codename,'w')
    iplen=len(baseaddr)
    reglen=len(regstruct)
    print('reglen: ', reglen)
    print(iplen,'iplen','baseaddr',baseaddr)
    with open ('demo_case.c','r', encoding='gbk', errors='ignore') as orcfile:
        for line in orcfile:
            print(line,file=wfid)
            if 'uint32_t module_baseaddr[] =' in line:
                i = 1
                for basad in baseaddr:
                    if i != iplen:
                        print(basad,',\n', file=wfid)
                    else:
                        print(basad,'\n', file=wfid)
                    i = i + 1
            if 'RegTypeDef IP_Reg[] =' in line:
                i = 1
                print('{', file=wfid)
                # print(type(regstruct))
                for reg in regstruct:
                    if i != reglen:
                        print('\t{', reg[0], reg[1], reg[2], '},\n', file=wfid)
                    else:
                        print('\t{', reg[0], reg[1], reg[2], '}\n', file=wfid)
                    i = i + 1
                print('};', file=wfid)
            if 'switch(module_baseaddr[i])' in line:
                i = 1
                print('         {', file=wfid)
                for basad in baseaddr:
                    if i!=iplen:
                        print ('          case',basad,':\n           Test_CAN_WriteReg(IP_Reg,',basad,', num); \n           break; \n',file=wfid)
                    else:
                        print ('          case',basad,':\n           Test_CAN_WriteReg(IP_Reg,',basad,', num); \n           break; \n',file=wfid)
                        print ('          default: printf("No module reg need to check!");',file=wfid)
                    i = i + 1
                print('         }', file=wfid)
    wfid.close()

# 将每个寄存器的offset、DefaultValue和WriteMask填入到regstruct中
def csv2reg (filename):
    print("start csv2reg...")
    regstruct=[]
    with open (filename,'r', encoding='gbk', errors='ignore') as fileid:
        reader=csv.reader(fileid)
        lst=list(reader)
        linenum=0
        for templist in lst:
            linenum=linenum+1
            if "offset" in templist:
                regAd   = linenum
                offad   = templist.index("offset")
                widthad =templist.index("width")
                bitad   = templist.index('bit')
                accad   =templist.index('Access')
                rstad   =templist.index('Reset')
                reglst=lst[regAd:]
                print('csv2reg')
                break
        regnum=0
        regline=[]
        rownum=0
        for row in reglst:
            print('in row for')
            rownum=rownum+1
            print(rownum)
            bits=row[bitad]
            bits=bits[1:-1]
            bitlist=bits.split(':')
            print(bitlist)
            if len(bitlist)==1:
                print(bitlist[0])
                lowad=int(bitlist[0])
                length=1
            else:
                i=int(bitlist[1])
                j=int(bitlist[-2])
                if i>j:
                     lowad=j
                     length=i-j
                else:
                    lowad=i
                    length=j-i
                length=length+1
            if row[regAd]!='':
                if regnum!=0:   #not the first register
                    mask=hex(mask)
                    regline.append(mask)
                    print(regline)
                    regstruct.append(regline)
                    #print the sturct
                else:
                    regnum=1
                mask=0
                default=[row[rstad]]
                regstr=[row[regAd]]
                width=row[widthad]
                #bits=row[bitad]
                #update the regline
                regline=[row[offad],row[rstad]]
                print(bits)
            else:
                pass
            if row[accad]!='RW':
                mask=mask
            else:
                data=(2**length-1)<<lowad
                pdata=hex(data)
                print(pdata)
                print('!!!!!!!!!!')
                mask=mask+data
            data1=row[rstad]
            #data=int(data1,base=16)
            #print(data)
            #default=default+data
            print('mask is ')
            print( mask)
            print('defaut is ')
            print(default)
            hmask=hex(mask)
        regline.append(hmask)
        print(regline)
        # print(type(mask))
        regstruct.append(regline)
        print('regstruct: ', regstruct)
        return regstruct

if __name__ == '__main__':
    word2reg()
    input('Press <Enter> to exit ......')

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
